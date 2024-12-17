import io
import os
import tempfile
import zipfile

from django.http import HttpResponse
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

from .models import ObjectActs


def compose_cert_file(obj: ObjectActs):
    download_file = tempfile.NamedTemporaryFile(prefix='certificates-')
    download_file_name = download_file.name
    download_file.close()
    result_doc = Document()
    for act in obj.acts.all():
        if act.certificates.count() > 0:
            for cert in act.certificates.all():
                with zipfile.ZipFile(cert.filename) as zf:
                    for name in zf.infolist():
                        if name.filename.startswith('word/media/'):
                            with zf.open(name.filename) as image:
                                paragraph = result_doc.add_paragraph(
                                    f'Приложение к АОСР №{obj.acts_prefix}-{act.act_number} от {act.act_date}')
                                paragraph_format = paragraph.paragraph_format
                                paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                                result_doc.add_picture(image, width=Cm(16.5), height=Cm(20.9))
                                result_doc.add_page_break()
    result_doc.save(download_file_name)
    with open(download_file_name, 'rb') as download_file:
        buffer = io.BytesIO()
        buffer.write(download_file.read())
        buffer.seek(0)
    if os.path.exists(download_file_name):
        os.remove(download_file_name)
    return buffer

