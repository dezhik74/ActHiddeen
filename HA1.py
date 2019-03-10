# -*- coding: utf-8 -*-

from docxtpl import DocxTemplate
from docx import Document
import json

def append_to_doc(doc,fname):
    t = Document(fname)
    for p in t.paragraphs:
        doc.add_paragraph("",p.style)       # add an empty paragraph in the matching style
        for r in p.runs:
            nr = doc.paragraphs[-1].add_run(r.text)
            nr.bold = r.bold
            nr.italic = r.italic
            nr.underline = r.underline

def Main():
    doc = DocxTemplate("HiddenActTemtplate1.docx")



    context = { 'Address' : "Губина ул.,  д.14 литера А",
                'SystemType' : "горячего водоснабжения",
                'Designer' : "ООО «ИЦ «ЛИФТ-ДИАГНОСТИКА» 192171, г. Санкт-Петербург, ул. Бабушкина, д. 36, к. 1, пом. 220, т./ф. 320-55-80/ 336-74-49",
                'ActNumber': "ТС-2",
                'ActDate': '"___"____________ 20__ г.',
                'SupervisionEngineerFull' : "Инженер ОСК№3 Гуц Н.Н. (ответственный по приказу №110 от 14.02.2018)",
                'ContractorEngineerFull' : "Руководитель проекта Узун П.Д.., (ответственный по приказу №18/2-охр от 28.02.2018)",
                'PresentedWork' : "Монтаж гильз в стенах и перекрытиях для прокладки трубопроводов системы отопления",
                'Project' : "проект №5-863/Г/ТС/2019-9-ОВ",
                'Materials' : "Трубы стальные гост 3262-75 ду32-ду100 (Сертификат РОСС RU.АГ99.Н06368 от 23.03.2016)",
                'Documents': "исполнительная схема системы ТС",
                'BeginDate': '"___"____________ 20__ г.',
                'EndDate': '"___"____________ 20__ г.',
                'PermittedWorks': "Прокладка трубопроводов системы отопления",
                'SupervisionEngineerShort': "Инженер ОСК№3 Гуц Н.Н.",
                'ContractorEngineerShort': "Руководитель проекта Узун П.Д."
                }

     # print(json.dumps(context, indent = 4))
     #
     # with open("123.json", "w", encoding="utf-8") as file:
     #     json.dump(context, file, ensure_ascii=False, separators=(",\n", ": "))
     #     file.close()
     #
     # with open('123.json', encoding="utf-8") as f:
     #     data = json.load(f)
     #
     # if context == data:
     #     print ('похожи')
     #
     # print(json.dumps(data, indent=4))

    doc.render(context)

    doc.save("generated.docx")

def doc_append (file1, file2):
    doc1 = Document(file1)
    doc2 = Document(file2)
    doc1.add_page_break()
    for el in doc2.element.body:
        doc1.element.body.append(el)
    doc1.save


def jsontest():
    with open('package.json', encoding="utf-8") as f:
        data = json.load(f)
    print (json.dumps(data, indent=4))



if __name__ =='__main__':
 #   jsontest()
    doc_append('generated1.docx','generated2.docx')
    Main()
